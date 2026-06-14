"""Export routines for the daily report worker."""

from __future__ import annotations

import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

PROJECT_ROOT = Path(__file__).resolve().parents[3]
TEMPLATES_ROOT = Path(__file__).resolve().parent / "templates"
RUNTIME_ROOT = PROJECT_ROOT / ".data" / "daily-report" / "runtime"
ARTIFACT_ROOT = RUNTIME_ROOT / "artifacts"

ONE_PAGE_CHAR_BUDGET = 1000
DOCX_TEMPLATE_NAME = "日报模板.docx"
DOCX_NAMING_PATTERN = re.compile(r"^国际日报-(\d{8})-(\d{3,})\.docx$")
XLSX_NAMING_PATTERN = re.compile(r"^resource-pool-\d{8}\.xlsx$")


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _parse_selections(raw_selections: list[str]) -> list[dict]:
    selections = []
    for raw in raw_selections:
        try:
            selections.append(json.loads(raw))
        except json.JSONDecodeError as error:
            raise SystemExit(f"Invalid selection JSON: {error}") from error
    return selections


def _resolve_artifact_path(object_key: str) -> Path:
    target = ARTIFACT_ROOT / object_key
    target.parent.mkdir(parents=True, exist_ok=True)
    return target


CHINESE_NUMERALS = ["一", "二", "三", "四", "五", "六"]


def _format_issue_date(issue_date: str) -> str:
    parsed = datetime.strptime(issue_date, "%Y-%m-%d")
    return f"{parsed.year}年{parsed.month}月{parsed.day}日"


def _set_paragraph_text_preserving_first_run(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _normalize_section_body(section: dict) -> str:
    body = str(section.get("body", "")).strip()
    title = str(section.get("title", "")).strip()
    if not body:
        return title
    # Avoid duplicated numbering when the model already returns "1. ..." / "一、...".
    body = re.sub(r"^\s*(?:\d+|[一二三四五六])\s*[\.、]\s*", "", body)
    return body


def _render_docx_from_template(*, target_path: Path, issue_date: str, issue_number: int, sections: list[dict]) -> bool:
    template_path = TEMPLATES_ROOT / DOCX_TEMPLATE_NAME
    if not template_path.exists():
        return False

    document = Document(template_path)
    non_empty_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if len(non_empty_paragraphs) < 7:
        return False

    header_text = f"第{issue_number}期                            {_format_issue_date(issue_date)}"
    _set_paragraph_text_preserving_first_run(non_empty_paragraphs[0], header_text)

    for index in range(6):
        body = _normalize_section_body(sections[index] if index < len(sections) else {})
        text = f"{CHINESE_NUMERALS[index]}、{body}"
        _set_paragraph_text_preserving_first_run(non_empty_paragraphs[index + 1], text)

    document.save(target_path)
    return True


def _render_docx(*, target_path: Path, issue_date: str, issue_number: int, sections: list[dict]) -> None:
    if _render_docx_from_template(
        target_path=target_path,
        issue_date=issue_date,
        issue_number=issue_number,
        sections=sections,
    ):
        return

    document = Document()

    heading = document.add_heading("国际日报", level=0)
    for run in heading.runs:
        run.font.size = Pt(20)

    subheader = document.add_paragraph()
    subheader.add_run(f"第{issue_number}期                            {_format_issue_date(issue_date)}").bold = True

    document.add_paragraph()

    for index, section in enumerate(sections[:6], start=1):
        body = _normalize_section_body(section)
        paragraph = document.add_paragraph(f"{CHINESE_NUMERALS[index - 1]}、{body}")
        paragraph.paragraph_format.first_line_indent = Pt(24)

    document.save(target_path)


def _measure_total_chars(sections: list[dict]) -> int:
    total = 0
    for section in sections:
        total += len(section.get("title", ""))
        total += len(section.get("body", ""))
    return total


def _append_resource_pool(
    *,
    template_path: Path,
    target_path: Path,
    issue_date: str,
    issue_number: int,
    selections: list[dict],
) -> None:
    if template_path.exists():
        # Read template + append
        from shutil import copyfile

        copyfile(template_path, target_path)
        workbook = load_workbook(target_path)
        worksheet = workbook.active
    else:
        # No template — create fresh
        from openpyxl import Workbook

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "资源池"
        worksheet.append(["期号", "出版日期", "序号", "标题", "来源", "源链接", "发布时间", "摘要"])

    for index, selection in enumerate(selections, start=1):
        worksheet.append([
            issue_number,
            issue_date,
            index,
            selection.get("title", ""),
            selection.get("sourceName", ""),
            selection.get("sourceUrl", ""),
            selection.get("publishedAt", ""),
            selection.get("summary", ""),
        ])

    target_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(target_path)


def _validate_export(
    *,
    docx_target: Path,
    xlsx_target: Path,
    issue_date: str,
    issue_number: int,
    expected_issue_number: int,
    sections: list[dict],
) -> dict:
    checks = []
    passed_all = True

    # Naming check: DOCX
    docx_match = DOCX_NAMING_PATTERN.match(docx_target.name)
    docx_naming_ok = bool(docx_match) and docx_match.group(1) == issue_date.replace("-", "")
    checks.append(
        {
            "code": "naming_rule",
            "scope": "docx_report",
            "passed": docx_naming_ok,
            "detail": {"fileName": docx_target.name, "expectedPattern": DOCX_NAMING_PATTERN.pattern},
        }
    )
    if not docx_naming_ok:
        passed_all = False

    # Naming check: XLSX
    xlsx_naming_ok = bool(XLSX_NAMING_PATTERN.match(xlsx_target.name))
    checks.append(
        {
            "code": "naming_rule",
            "scope": "resource_pool_xlsx",
            "passed": xlsx_naming_ok,
            "detail": {"fileName": xlsx_target.name, "expectedPattern": XLSX_NAMING_PATTERN.pattern},
        }
    )
    if not xlsx_naming_ok:
        passed_all = False

    # Issue-number check
    issue_number_ok = isinstance(expected_issue_number, int) and expected_issue_number > 0 and expected_issue_number == issue_number
    checks.append(
        {
            "code": "issue_number_match",
            "scope": "task",
            "passed": issue_number_ok,
            "detail": {"recorded": expected_issue_number, "embedded": issue_number},
        }
    )
    if not issue_number_ok:
        passed_all = False

    # One-page constraint check
    total_chars = _measure_total_chars(sections)
    one_page_ok = total_chars <= ONE_PAGE_CHAR_BUDGET
    checks.append(
        {
            "code": "one_page_budget",
            "scope": "docx_report",
            "passed": one_page_ok,
            "detail": {"totalChars": total_chars, "budget": ONE_PAGE_CHAR_BUDGET},
        }
    )
    if not one_page_ok:
        passed_all = False

    return {"passed": passed_all, "checks": checks}


def execute_export_command(args) -> dict:
    issue_date = args.issue_date
    issue_number = int(args.issue_number)
    sections = json.loads(args.sections_json)
    selections = _parse_selections(args.selection)

    docx_target = _resolve_artifact_path(args.docx_object_key)
    xlsx_target = _resolve_artifact_path(args.xlsx_object_key)

    try:
        _render_docx(
            target_path=docx_target,
            issue_date=issue_date,
            issue_number=issue_number,
            sections=sections,
        )
        _append_resource_pool(
            template_path=TEMPLATES_ROOT / "resource-pool.xlsx",
            target_path=xlsx_target,
            issue_date=issue_date,
            issue_number=issue_number,
            selections=selections,
        )
    except Exception as error:  # pragma: no cover - defensive
        return {
            "ok": False,
            "code": "export_render_failed",
            "message": f"渲染产物失败：{error}",
            "failureCategory": "worker_failure",
            "retriable": True,
            "aiCalls": [],
            "events": [],
        }

    validation_report = _validate_export(
        docx_target=docx_target,
        xlsx_target=xlsx_target,
        issue_date=issue_date,
        issue_number=issue_number,
        expected_issue_number=issue_number,
        sections=sections,
    )

    if not validation_report["passed"]:
        # Remove emitted files so a half-exported attempt does not pollute storage
        for target in (docx_target, xlsx_target):
            try:
                target.unlink()
            except FileNotFoundError:
                pass
        return {
            "ok": False,
            "code": "export_validation_failed",
            "message": "导出未通过校验。",
            "failureCategory": "validation_failure",
            "retriable": False,
            "validationReport": validation_report,
            "aiCalls": [],
            "events": [
                {
                    "type": "export_validation_failed",
                    "createdAt": _utc_now(),
                    "detail": {"checkCount": len(validation_report["checks"])},
                }
            ],
        }

    return {
        "ok": True,
        "docx": {
            "fileName": docx_target.name,
            "sizeBytes": docx_target.stat().st_size,
        },
        "xlsx": {
            "fileName": xlsx_target.name,
            "sizeBytes": xlsx_target.stat().st_size,
        },
        "validationReport": validation_report,
        "summary": {
            "draftGenerated": True,
            "exportCompleted": True,
            "sectionCount": len(sections),
            "selectionCount": len(selections),
        },
        "aiCalls": [],
        "events": [
            {
                "type": "export_completed",
                "createdAt": _utc_now(),
                "detail": {"docxSizeBytes": docx_target.stat().st_size, "xlsxSizeBytes": xlsx_target.stat().st_size},
            }
        ],
    }